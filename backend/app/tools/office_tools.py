"""
OFİS ve ARŞİV ARAÇLARI
ZIP, PDF, Word, Excel yönetimi
"""
from __future__ import annotations

import json
import os
import subprocess
import zipfile
import tarfile
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from ..config import settings


def _resolve_path(path: str) -> Path:
    if not path or path == ".":
        return settings.workspace_path.resolve()
    p = Path(path).expanduser()
    if p.is_absolute():
        return p.resolve()
    return (settings.workspace_path / p).resolve()


def _assert_within(base: Path, target: Path) -> None:
    base_resolved = base.resolve()
    target_resolved = target.resolve()
    if target_resolved != base_resolved and base_resolved not in target_resolved.parents:
        raise ValueError("Arsiv ici yol traversal denemesi engellendi.")


# =============================================================================
# ZIP / ARŞİV ARAÇLARI
# =============================================================================

def tool_create_zip(source_path: str, output_path: str = "", password: str = "") -> Dict[str, Any]:
    """ZIP arşivi oluştur."""
    try:
        source = _resolve_path(source_path)
        
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"{source.name}_{timestamp}.zip"
        
        target = _resolve_path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        
        if password:
            # Şifreli ZIP (pyminizip veya benzeri gerekebilir)
            # Şimdilik standart ZIP
            pass
        
        with zipfile.ZipFile(target, 'w', zipfile.ZIP_DEFLATED) as zipf:
            if source.is_file():
                zipf.write(source, source.name)
            elif source.is_dir():
                for root, dirs, files in os.walk(source):
                    for file in files:
                        file_path = Path(root) / file
                        arcname = file_path.relative_to(source)
                        zipf.write(file_path, arcname)
        
        return {
            "success": True,
            "source": str(source),
            "archive": str(target),
            "size": target.stat().st_size
        }
    except Exception as e:
        return {"error": str(e)}


def tool_extract_zip(zip_path: str, output_dir: str = "", password: str = "") -> Dict[str, Any]:
    """ZIP arşivini çıkar."""
    try:
        zip_file = _resolve_path(zip_path)
        
        if not output_dir:
            output_dir = str(zip_file.parent / zip_file.stem)
        
        target_dir = _resolve_path(output_dir)
        target_dir.mkdir(parents=True, exist_ok=True)
        
        with zipfile.ZipFile(zip_file, 'r') as zipf:
            file_list = zipf.namelist()
            for info in zipf.infolist():
                member_path = (target_dir / info.filename).resolve()
                _assert_within(target_dir, member_path)
                if info.is_dir():
                    member_path.mkdir(parents=True, exist_ok=True)
                    continue
                member_path.parent.mkdir(parents=True, exist_ok=True)
                with zipf.open(info, pwd=password.encode() if password else None) as src, open(member_path, "wb") as dst:
                    shutil.copyfileobj(src, dst)
        
        return {
            "success": True,
            "archive": str(zip_file),
            "extracted_to": str(target_dir),
            "file_count": len(file_list),
            "files": file_list[:50]  # İlk 50 dosya
        }
    except Exception as e:
        return {"error": str(e)}


def tool_list_zip_contents(zip_path: str) -> Dict[str, Any]:
    """ZIP içeriğini listele."""
    try:
        zip_file = _resolve_path(zip_path)
        
        with zipfile.ZipFile(zip_file, 'r') as zipf:
            files = []
            total_size = 0
            
            for info in zipf.infolist():
                files.append({
                    "name": info.filename,
                    "size": info.file_size,
                    "compressed": info.compress_size,
                    "date": datetime(*info.date_time).isoformat()
                })
                total_size += info.file_size
        
        return {
            "archive": str(zip_file),
            "file_count": len(files),
            "total_size": total_size,
            "files": files[:100]
        }
    except Exception as e:
        return {"error": str(e)}


def tool_create_tar(source_path: str, output_path: str = "", compression: str = "gz") -> Dict[str, Any]:
    """TAR arşivi oluştur (.tar.gz, .tar.bz2, .tar.xz)."""
    try:
        source = _resolve_path(source_path)
        
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            ext = {"gz": "tar.gz", "bz2": "tar.bz2", "xz": "tar.xz"}.get(compression, "tar")
            output_path = f"{source.name}_{timestamp}.{ext}"
        
        target = _resolve_path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        
        mode = {"gz": "w:gz", "bz2": "w:bz2", "xz": "w:xz"}.get(compression, "w")
        
        with tarfile.open(target, mode) as tar:
            tar.add(source, arcname=source.name)
        
        return {
            "success": True,
            "source": str(source),
            "archive": str(target),
            "compression": compression,
            "size": target.stat().st_size
        }
    except Exception as e:
        return {"error": str(e)}


def tool_extract_tar(tar_path: str, output_dir: str = "") -> Dict[str, Any]:
    """TAR arşivini çıkar."""
    try:
        tar_file = _resolve_path(tar_path)
        
        if not output_dir:
            output_dir = str(tar_file.parent / tar_file.stem)
        
        target_dir = _resolve_path(output_dir)
        target_dir.mkdir(parents=True, exist_ok=True)
        
        with tarfile.open(tar_file, 'r:*') as tar:
            members = tar.getmembers()
            for member in members:
                member_path = (target_dir / member.name).resolve()
                _assert_within(target_dir, member_path)
            tar.extractall(target_dir)
        
        return {
            "success": True,
            "archive": str(tar_file),
            "extracted_to": str(target_dir),
            "file_count": len(members)
        }
    except Exception as e:
        return {"error": str(e)}


# =============================================================================
# PDF ARAÇLARI
# =============================================================================

def tool_read_pdf(pdf_path: str, page_start: int = 0, page_end: int = None) -> Dict[str, Any]:
    """PDF dosyasını oku."""
    try:
        import PyPDF2
        
        pdf_file = _resolve_path(pdf_path)
        
        with open(pdf_file, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            num_pages = len(reader.pages)
            
            page_end = page_end or min(page_start + 5, num_pages)
            text = ""
            
            for i in range(page_start, min(page_end, num_pages)):
                page = reader.pages[i]
                text += f"\n--- Sayfa {i+1} ---\n"
                text += page.extract_text()
        
        return {
            "path": str(pdf_file),
            "total_pages": num_pages,
            "read_pages": f"{page_start+1}-{min(page_end, num_pages)}",
            "text": text[:15000]  # Limit
        }
    except Exception as e:
        return {"error": str(e)}


def tool_create_pdf(output_path: str, title: str = "", content: str = "") -> Dict[str, Any]:
    """Türkçe destekli, düzgün formatlı PDF oluştur."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        target = _resolve_path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)

        c = canvas.Canvas(str(target), pagesize=A4)
        width, height = A4
        margin_left = 2 * cm
        margin_right = 2 * cm
        margin_top = 2.5 * cm
        margin_bottom = 2 * cm
        usable_width = width - margin_left - margin_right

        # Türkçe karakter desteği - Arial kullan, yoksa fallback
        font_name = "Helvetica"
        font_name_bold = "Helvetica-Bold"
        for font_path in [
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/Arial.ttf",
            "C:/Windows/Fonts/calibri.ttf",
        ]:
            try:
                pdfmetrics.registerFont(TTFont("TRFont", font_path))
                bold_path = font_path.replace("arial", "arialbd").replace("calibri", "calibrib")
                try:
                    pdfmetrics.registerFont(TTFont("TRFontBold", bold_path))
                    font_name_bold = "TRFontBold"
                except Exception:
                    font_name_bold = "TRFont"
                font_name = "TRFont"
                break
            except Exception:
                continue

        def _wrap_text(text: str, font: str, size: int, max_w: float) -> List[str]:
            """Metni satır genişliğine göre sar."""
            words = text.split()
            if not words:
                return [""]
            wrapped: List[str] = []
            current = words[0]
            for word in words[1:]:
                test = current + " " + word
                tw = c.stringWidth(test, font, size)
                if tw <= max_w:
                    current = test
                else:
                    wrapped.append(current)
                    current = word
            wrapped.append(current)
            return wrapped

        page_count = 1
        y = height - margin_top

        # Başlık
        if title:
            c.setFont(font_name_bold, 16)
            title_lines = _wrap_text(title, font_name_bold, 16, usable_width)
            for tl in title_lines:
                if y < margin_bottom:
                    c.showPage()
                    page_count += 1
                    y = height - margin_top
                c.setFont(font_name_bold, 16)
                c.drawString(margin_left, y, tl)
                y -= 22
            # Başlık altı boşluk
            y -= 10
            # Alt çizgi
            c.setStrokeColorRGB(0.7, 0.7, 0.7)
            c.line(margin_left, y + 5, width - margin_right, y + 5)
            y -= 15

        # İçerik
        font_size = 11
        line_height = 15
        c.setFont(font_name, font_size)
        paragraphs = content.split('\n')

        for para in paragraphs:
            para = para.rstrip()
            if not para:
                y -= line_height * 0.6
                if y < margin_bottom:
                    c.showPage()
                    page_count += 1
                    y = height - margin_top
                    c.setFont(font_name, font_size)
                continue

            # Markdown-tarzı başlık desteği
            is_heading = False
            if para.startswith("### "):
                c.setFont(font_name_bold, 12)
                para = para[4:]
                is_heading = True
            elif para.startswith("## "):
                c.setFont(font_name_bold, 13)
                para = para[3:]
                is_heading = True
                y -= 5
            elif para.startswith("# "):
                c.setFont(font_name_bold, 14)
                para = para[2:]
                is_heading = True
                y -= 8

            active_font = font_name_bold if is_heading else font_name
            active_size = 14 if (is_heading and para.startswith("# ")) else (13 if is_heading else font_size)

            wrapped = _wrap_text(para, active_font, active_size, usable_width)
            for wl in wrapped:
                if y < margin_bottom:
                    c.showPage()
                    page_count += 1
                    y = height - margin_top
                    c.setFont(active_font, active_size)
                c.drawString(margin_left, y, wl)
                y -= line_height

            if is_heading:
                y -= 4
                c.setFont(font_name, font_size)

        # Sayfa numaraları
        total = page_count
        for i in range(1, total + 1):
            if i > 1:
                # İlk sayfa zaten açık, diğerleri showPage ile oluşturuldu
                pass
        # Sayfa numarası ekleme (basit - son sayfaya)
        c.setFont(font_name, 8)
        c.drawCentredString(width / 2, margin_bottom - 15,
                            f"Sayfa {page_count} / {page_count}")

        c.save()

        return {
            "success": True,
            "path": str(target),
            "title": title,
            "pages": page_count,
            "language": "Türkçe"
        }
    except Exception as e:
        return {"error": str(e)}


def tool_merge_pdfs(pdf_list: List[str], output_path: str) -> Dict[str, Any]:
    """Birden fazla PDF'i birleştir."""
    try:
        import PyPDF2
        
        merger = PyPDF2.PdfMerger()
        
        for pdf_path in pdf_list:
            pdf_file = _resolve_path(pdf_path)
            merger.append(str(pdf_file))
        
        target = _resolve_path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        
        merger.write(str(target))
        merger.close()
        
        return {
            "success": True,
            "output": str(target),
            "merged_files": len(pdf_list)
        }
    except Exception as e:
        return {"error": str(e)}


def tool_split_pdf(pdf_path: str, page_ranges: List[Dict], output_prefix: str = "") -> Dict[str, Any]:
    """PDF'i sayfa aralıklarına göre böl."""
    try:
        import PyPDF2
        
        pdf_file = _resolve_path(pdf_path)
        
        if not output_prefix:
            output_prefix = str(pdf_file.parent / f"{pdf_file.stem}_split")
        
        with open(pdf_file, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            output_files = []
            
            for i, range_info in enumerate(page_ranges):
                start = range_info.get("start", 0)
                end = range_info.get("end", len(reader.pages))
                name = range_info.get("name", f"part_{i+1}")
                
                writer = PyPDF2.PdfWriter()
                for page_num in range(start, min(end, len(reader.pages))):
                    writer.add_page(reader.pages[page_num])
                
                output_path = f"{output_prefix}_{name}.pdf"
                with open(output_path, 'wb') as out:
                    writer.write(out)
                
                output_files.append(output_path)
        
        return {
            "success": True,
            "output_files": output_files,
            "count": len(output_files)
        }
    except Exception as e:
        return {"error": str(e)}


# =============================================================================
# WORD (DOCX) ARAÇLARI
# =============================================================================

def tool_create_docx(output_path: str, title: str = "", paragraphs: List[str] = None, 
                     headings: List[Dict] = None, tables: List[List[List]] = None) -> Dict[str, Any]:
    """Word belgesi (.docx) oluştur."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Başlık
        if title:
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Başlıklar ve içerik
        if headings:
            for h in headings:
                text = h.get("text", "")
                level = h.get("level", 1)
                doc.add_heading(text, level=level)
                
                if "content" in h:
                    doc.add_paragraph(h["content"])
        
        # Paragraflar
        if paragraphs:
            for para in paragraphs:
                doc.add_paragraph(para)
        
        # Tablolar
        if tables:
            for table_data in tables:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = str(cell)
        
        # Kaydet
        target = _resolve_path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(target))
        
        return {
            "success": True,
            "path": str(target),
            "title": title,
            "paragraphs": len(paragraphs) if paragraphs else 0
        }
    except Exception as e:
        return {"error": str(e)}


def tool_read_docx(docx_path: str) -> Dict[str, Any]:
    """Word belgesini oku."""
    try:
        from docx import Document
        
        doc_file = _resolve_path(docx_path)
        doc = Document(str(doc_file))
        
        text = []
        for para in doc.paragraphs:
            if para.text:
                text.append(para.text)
        
        # Tablolar
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            "path": str(doc_file),
            "paragraphs": len(text),
            "text": "\n".join(text)[:10000],
            "tables": len(tables)
        }
    except Exception as e:
        return {"error": str(e)}


def tool_add_to_docx(docx_path: str, paragraphs: List[str] = None, 
                     heading: str = "", heading_level: int = 1) -> Dict[str, Any]:
    """Word belgesine ekleme yap."""
    try:
        from docx import Document
        
        doc_file = _resolve_path(docx_path)
        
        if doc_file.exists():
            doc = Document(str(doc_file))
        else:
            doc = Document()
        
        if heading:
            doc.add_heading(heading, level=heading_level)
        
        if paragraphs:
            for para in paragraphs:
                doc.add_paragraph(para)
        
        doc.save(str(doc_file))
        
        return {
            "success": True,
            "path": str(doc_file),
            "added_paragraphs": len(paragraphs) if paragraphs else 0
        }
    except Exception as e:
        return {"error": str(e)}


# =============================================================================
# EXCEL (XLSX) ARAÇLARI
# =============================================================================

def tool_create_excel(output_path: str, sheet_name: str = "Sheet1", 
                      data: List[List] = None, headers: List[str] = None) -> Dict[str, Any]:
    """Excel dosyası (.xlsx) oluştur."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
        
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name
        
        # Başlıklar
        if headers:
            ws.append(headers)
            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        
        # Veriler
        if data:
            for row in data:
                ws.append(row)
        
        # Kaydet
        target = _resolve_path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(target))
        
        return {
            "success": True,
            "path": str(target),
            "sheet": sheet_name,
            "rows": len(data) if data else 0,
            "columns": len(headers) if headers else 0
        }
    except Exception as e:
        return {"error": str(e)}


def tool_read_excel(excel_path: str, sheet_name: str = "", max_rows: int = 100) -> Dict[str, Any]:
    """Excel dosyasını oku."""
    try:
        from openpyxl import load_workbook
        
        excel_file = _resolve_path(excel_path)
        wb = load_workbook(str(excel_file), data_only=True)
        
        if sheet_name:
            ws = wb[sheet_name]
        else:
            ws = wb.active
            sheet_name = ws.title
        
        data = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                break
            data.append(row)
        
        return {
            "path": str(excel_file),
            "sheet": sheet_name,
            "total_rows": ws.max_row,
            "columns": ws.max_column,
            "data": data
        }
    except Exception as e:
        return {"error": str(e)}


def tool_add_to_excel(excel_path: str, data: List[List], sheet_name: str = "") -> Dict[str, Any]:
    """Excel dosyasına veri ekle."""
    try:
        from openpyxl import load_workbook
        
        excel_file = _resolve_path(excel_path)
        
        if excel_file.exists():
            wb = load_workbook(str(excel_file))
        else:
            from openpyxl import Workbook
            wb = Workbook()
        
        if sheet_name:
            if sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
            else:
                ws = wb.create_sheet(title=sheet_name)
        else:
            ws = wb.active
        
        for row in data:
            ws.append(row)
        
        wb.save(str(excel_file))
        
        return {
            "success": True,
            "path": str(excel_file),
            "added_rows": len(data)
        }
    except Exception as e:
        return {"error": str(e)}


# =============================================================================
# DİĞER OFİS ARAÇLARI
# =============================================================================

def tool_open_in_vscode(path: str, wait: bool = False) -> Dict[str, Any]:
    """Dosya veya klasörü VS Code'da aç."""
    try:
        target = _resolve_path(path)
        
        if not target.exists():
            return {"error": "Dosya veya klasör bulunamadı", "path": str(target)}
        
        # VS Code'u aç
        cmd = ["code", str(target)]
        if wait:
            cmd.append("--wait")
        
        subprocess.Popen(cmd, shell=False)
        
        return {
            "success": True,
            "opened": str(target),
            "editor": "VS Code"
        }
    except Exception as e:
        return {"error": str(e), "note": "VS Code PATH'e ekli olduğundan emin olun"}


def tool_open_folder(folder_path: str) -> Dict[str, Any]:
    """Klasörü Dosya Gezgini'nde aç."""
    try:
        target = _resolve_path(folder_path)
        
        if not target.exists():
            return {"error": "Klasör bulunamadı", "path": str(target)}
        
        if os.name == 'nt':  # Windows
            subprocess.Popen(["explorer", str(target)], shell=False)
        else:  # Linux/Mac
            subprocess.Popen(['xdg-open', str(target)])
        
        return {
            "success": True,
            "opened": str(target)
        }
    except Exception as e:
        return {"error": str(e)}


def tool_create_folder(folder_path: str) -> Dict[str, Any]:
    """Yeni klasör oluştur."""
    try:
        target = _resolve_path(folder_path)
        target.mkdir(parents=True, exist_ok=True)
        
        return {
            "success": True,
            "path": str(target),
            "created": True
        }
    except Exception as e:
        return {"error": str(e)}


def tool_analyze_project_code(project_path: str, output_format: str = "json") -> Dict[str, Any]:
    """Proje kodlarını analiz et ve raporla."""
    try:
        target = _resolve_path(project_path)
        
        if not target.is_dir():
            return {"error": "Proje klasörü değil", "path": str(target)}
        
        stats = {
            "total_files": 0,
            "total_lines": 0,
            "languages": {},
            "files": []
        }
        
        # Dil uzantıları
        lang_map = {
            ".py": "Python",
            ".js": "JavaScript",
            ".ts": "TypeScript",
            ".jsx": "React",
            ".tsx": "ReactTS",
            ".java": "Java",
            ".cpp": "C++",
            ".c": "C",
            ".h": "Header",
            ".go": "Go",
            ".rs": "Rust",
            ".rb": "Ruby",
            ".php": "PHP",
            ".swift": "Swift",
            ".kt": "Kotlin"
        }
        
        for root, dirs, files in os.walk(target):
            # node_modules, .venv gibi klasörleri atla
            dirs[:] = [d for d in dirs if d not in ['node_modules', '.venv', '__pycache__', '.git', 'dist', 'build']]
            
            for file in files:
                ext = Path(file).suffix.lower()
                if ext in lang_map:
                    file_path = Path(root) / file
                    try:
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            lines = f.readlines()
                            line_count = len(lines)
                            
                            lang = lang_map[ext]
                            stats["total_files"] += 1
                            stats["total_lines"] += line_count
                            
                            if lang not in stats["languages"]:
                                stats["languages"][lang] = {"files": 0, "lines": 0}
                            stats["languages"][lang]["files"] += 1
                            stats["languages"][lang]["lines"] += line_count
                            
                            stats["files"].append({
                                "path": str(file_path.relative_to(target)),
                                "language": lang,
                                "lines": line_count
                            })
                    except:
                        pass
        
        # Rapor oluştur
        if output_format == "markdown":
            report_lines = [
                f"# Proje Analiz Raporu",
                f"",
                f"**Proje:** {target.name}",
                f"**Tarih:** {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                f"",
                f"## Özet",
                f"- Toplam Dosya: {stats['total_files']}",
                f"- Toplam Satır: {stats['total_lines']}",
                f"",
                f"## Diller",
            ]
            
            for lang, data in sorted(stats["languages"].items(), key=lambda x: x[1]["lines"], reverse=True):
                report_lines.append(f"- **{lang}:** {data['files']} dosya, {data['lines']} satır")
            
            report_lines.extend(["", "## Dosyalar"])
            for f in sorted(stats["files"], key=lambda x: x["lines"], reverse=True)[:20]:
                report_lines.append(f"- `{f['path']}` ({f['lines']} satır)")
            
            report_text = "\n".join(report_lines)
            
            report_path = target / "PROJECT_ANALYSIS.md"
            report_path.write_text(report_text, encoding='utf-8')
            
            return {
                "success": True,
                "stats": stats,
                "report": str(report_path)
            }
        
        return {
            "success": True,
            "stats": stats
        }
        
    except Exception as e:
        return {"error": str(e)}
